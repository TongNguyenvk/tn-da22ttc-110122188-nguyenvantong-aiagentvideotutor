"""
Document Renderer - Tao file DOCX tu plan va screenshots
"""
import logging
from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Inches, Pt

# Import base_renderer without relative import
import sys
sys.path.insert(0, str(Path(__file__).parent))
from base_renderer import Renderer

logger = logging.getLogger(__name__)

class DocumentRenderer(Renderer):
    """Renderer de tao file DOCX"""
    
    def render(self, plan: Dict, artifacts: Dict) -> str:
        """Tao file DOCX tu plan va screenshots"""
        logger.info("  [DocumentRenderer] Bat dau render DOCX...")
        
        doc = Document()
        self._setup_styles(doc)
        
        # Them tieu de
        title = plan.get('title', 'Tutorial')
        doc.add_heading(title, level=0)
        
        # Them tung buoc
        steps = plan.get('steps', [])
        screenshots = artifacts.get('screenshots', [])
        
        for i, step in enumerate(steps):
            narration = step.get('narration', f'Buoc {i+1}')
            
            # Them tieu de buoc
            doc.add_heading(f'Buoc {i+1}: {narration}', level=1)
            
            # Them anh neu co
            if i < len(screenshots) and screenshots[i]:
                screenshot_path = screenshots[i]
                if Path(screenshot_path).exists():
                    try:
                        doc.add_picture(screenshot_path, width=Inches(6))
                    except Exception as e:
                        logger.error(f"  [DocumentRenderer] Loi chen anh: {e}")
            
            # Them khoang trang
            doc.add_paragraph()
        
        # Luu file
        output_path = self.output_dir / f"{plan.get('name', 'tutorial')}.docx"
        doc.save(output_path)
        
        logger.info(f"  [DocumentRenderer] Da luu DOCX: {output_path}")
        return str(output_path)
    
    def _setup_styles(self, doc):
        """Thiet lap style cho document"""
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    @property
    def output_format(self) -> str:
        return "docx"
